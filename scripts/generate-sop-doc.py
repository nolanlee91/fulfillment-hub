# -*- coding: utf-8 -*-
"""
Generate Word document (SOP) for KDExpress Fulfillment App workflow.
Reads images from 'SOP app/' folder and embeds them into a structured .docx.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- paths ---
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SOP_DIR = os.path.join(ROOT, "SOP app")
OUT_PATH = os.path.join(ROOT, "SOP_KDExpress_Fulfillment.docx")

IMG = {
    "flow": os.path.join(SOP_DIR, "1. Luồng follow.png"),
    "sync": os.path.join(SOP_DIR, "2. Sync.png"),
    "ready": os.path.join(SOP_DIR, "3. Ready.png"),
    "bulk": os.path.join(SOP_DIR, "4. Bulk action.png"),
    "batch": os.path.join(SOP_DIR, "5. Tạo batch.png"),
    "download": os.path.join(SOP_DIR, "6. Download file.png"),
    "upload": os.path.join(SOP_DIR, "7. Upload labeled.png"),
    "error": os.path.join(SOP_DIR, "8. Báo khách lỗi.png"),
}


# --- helpers ---
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)  # KDExpress green
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullet(doc, text, bold_first=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_first:
        r1 = p.add_run(bold_first)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p


def add_image_centered(doc, path, width_inches=6.3, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def add_status_badge_line(doc, color_hex, label, desc):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("● ")
    r.font.color.rgb = RGBColor(*tuple(int(color_hex[i : i + 2], 16) for i in (0, 2, 4)))
    r.bold = True
    r.font.size = Pt(12)
    r2 = p.add_run(label + ": ")
    r2.bold = True
    r2.font.size = Pt(11)
    r3 = p.add_run(desc)
    r3.font.size = Pt(11)


# --- build document ---
doc = Document()

# Margins
for section in doc.sections:
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

# Title
title = doc.add_heading("SOP — Quy trình xử lý đơn hàng trên KDExpress Fulfillment Hub", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Tài liệu hướng dẫn vận hành — fulfillment.hub")
sr.italic = True
sr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
sr.font.size = Pt(11)

doc.add_paragraph()

# =================================================================
# I. TỔNG QUAN LUỒNG FOLLOW
# =================================================================
add_heading(doc, "I. Tổng quan luồng follow đơn hàng", level=1)

add_para(
    doc,
    "Mỗi đơn hàng đi qua một pipeline cố định gồm 8 bước. App tự động cập nhật trạng thái "
    "(status) khi đơn chuyển từ bước này sang bước khác. Status pipeline được hiển thị trực "
    "quan ở Dashboard.",
)

add_image_centered(doc, IMG["flow"], width_inches=6.5, caption="Hình 1 — Order Processing Flow (Status Pipeline)")

add_heading(doc, "Sơ đồ luồng", level=2)
add_para(
    doc,
    "Customer Google Sheets  →  SYNC  →  READY  →  CREATE BATCH  →  CLICKSHIP/EST "
    "(uploaded awaiting label)  →  DOWNLOAD LABEL  →  SAVE LABEL TO DEVICE  →  "
    "UPLOAD LABEL TO KDE-APP (shipping reconciliation)  →  LABELED / IN TRANSIT / NEEDS ATTENTION",
    italic=True,
)

add_heading(doc, "Ý nghĩa từng trạng thái", level=2)
add_status_badge_line(doc, "3B82F6", "New", "Đơn vừa được sync về từ Google Sheets, chưa validate.")
add_status_badge_line(doc, "16A34A", "Ready", "Đơn đã pass validate (đủ name, address, phone hợp lệ), sẵn sàng đóng gói + tạo label.")
add_status_badge_line(doc, "DC2626", "Error", "Đơn thiếu data (Name/Address/Phone) hoặc ClickShip/EST từ chối tạo label. Cần báo khách sửa.")
add_status_badge_line(doc, "F59E0B", "Updated", "Đơn ERROR trước đây, khách đã sửa data trong sheet, sync về thấy có thay đổi → cần validate lại.")
add_status_badge_line(doc, "6B7280", "Exported", "Đơn đã nằm trong batch + đã xuất file Excel up lên ClickShip/EST, đang chờ label.")
add_status_badge_line(doc, "8B5CF6", "Label", "Đơn đã có tracking number, label đã được tạo thành công.")
add_status_badge_line(doc, "0EA5E9", "In Transit", "Đơn đang trên đường giao, có tracking event từ Canada Post/GLS.")

doc.add_page_break()

# =================================================================
# II. CHI TIẾT TỪNG BƯỚC
# =================================================================
add_heading(doc, "II. Chi tiết từng bước vận hành", level=1)

# --- STEP 1: SYNC ---
add_heading(doc, "Bước 1 — SYNC: Kéo dữ liệu đơn mới từ Google Sheets", level=2)

add_para(doc, "Mục đích:", bold=True)
add_para(
    doc,
    "App đọc data từ các sheet nguồn của khách (Venatureco, Skylane, CosmeticsTuanlx, ...), "
    "tạo đơn mới trong DB. Đơn nào chưa có tracking number trong sheet thì mới được pull.",
)

add_para(doc, "Thao tác:", bold=True)
add_bullet(doc, "Vào Dashboard.")
add_bullet(doc, 'Bấm nút "Sync" ở góc phải trên cùng.')
add_bullet(doc, "App pull data → tạo đơn NEW → tự động validate → chuyển sang READY (nếu OK) hoặc ERROR (nếu thiếu data).")

add_image_centered(doc, IMG["sync"], width_inches=6.5, caption="Hình 2 — Dashboard trước khi sync (0 READY)")

add_para(doc, "Lưu ý:", bold=True)
add_bullet(doc, "Sync chỉ kéo đơn MỚI (chưa có trong DB) hoặc đơn ERROR/ERROR_UPDATED.", bold_first="")
add_bullet(doc, "Đơn đã EXPORTED/LABEL_CREATED/IN_TRANSIT sẽ KHÔNG bị overwrite nếu khách sửa lại trong sheet.")
add_bullet(doc, "Sync chạy ~5-30 giây tùy số sheet.")

doc.add_page_break()

# --- STEP 2: READY ---
add_heading(doc, "Bước 2 — READY: Kiểm tra số đơn sẵn sàng đóng gói", level=2)

add_para(doc, "Mục đích:", bold=True)
add_para(
    doc,
    "Sau khi sync xong, Dashboard hiển thị số đơn READY (đã pass validate). Đây là số đơn "
    "thực tế có thể đóng gói + tạo label ngay.",
)

add_image_centered(doc, IMG["ready"], width_inches=6.5, caption="Hình 3 — Dashboard sau khi sync (196 READY)")

add_para(doc, "Quan sát các card chỉ số:", bold=True)
add_bullet(doc, "IN PROGRESS: tổng đơn đang xử lý (NEW + READY + ERROR + EXPORTED + LABEL_CREATED + IN_TRANSIT).")
add_bullet(doc, "READY TO PACK: số đơn READY, sẵn sàng đưa vào batch.")
add_bullet(doc, "DELIVERED: đơn đã giao thành công.")
add_bullet(doc, "NEEDS ATTENTION: đơn cần chú ý (notice card, delay, stuck) — kiểm tra ngay.")
add_bullet(doc, "RED/YELLOW FLAG ORDERS: đơn được gắn cờ trao đổi với khách.")

doc.add_page_break()

# --- STEP 3: BULK ACTION ---
add_heading(doc, "Bước 3 — BULK ACTION: Chọn đơn READY để tạo batch", level=2)

add_para(doc, "Mục đích:", bold=True)
add_para(
    doc,
    "Vào trang Active Orders, filter status = Ready, tick chọn các đơn cần đưa vào lô đóng gói. "
    "Có thể chọn tất cả, hoặc chỉ chọn 1 customer/product cụ thể.",
)

add_para(doc, "Thao tác:", bold=True)
add_bullet(doc, "Sidebar → Active Orders.")
add_bullet(doc, 'Filter (chip tab phía trên): bấm "Ready" để chỉ hiện đơn READY.')
add_bullet(doc, "Tùy chọn lọc thêm theo Customer / Product / Payment / Search.")
add_bullet(doc, "Tick checkbox bên trái mỗi đơn, hoặc tick checkbox header để chọn tất cả.")
add_bullet(doc, 'Khi đã chọn → 3 nút hành động sẽ hiện ở góc phải: Export / Delete / Create Batch.')

add_image_centered(doc, IMG["bulk"], width_inches=6.5, caption="Hình 4 — Active Orders, filter Ready, sẵn sàng bulk action")

doc.add_page_break()

# --- STEP 4: CREATE BATCH ---
add_heading(doc, "Bước 4 — CREATE BATCH: Tạo lô đóng gói", level=2)

add_para(doc, "Mục đích:", bold=True)
add_para(
    doc,
    'Gom các đơn READY đã chọn thành 1 "batch" (lô) duy nhất. Batch là đơn vị để xuất file '
    "Excel up lên ClickShip/EST. Mỗi batch có ID format: YYYY-MM-DD-AM/PM-NNN "
    "(vd 2026-05-27-PM-003).",
)

add_para(doc, "Thao tác:", bold=True)
add_bullet(doc, "Sau khi đã tick chọn các đơn READY ở bước 3.")
add_bullet(doc, 'Bấm nút "Create Batch" góc phải.')
add_bullet(doc, "App tạo batch mới, gắn batchId vào các đơn, chuyển status: READY → EXPORTED.")
add_bullet(doc, "Đơn EXPORTED giờ đã được khóa, không sửa data từ sheet nữa được.")

add_image_centered(doc, IMG["batch"], width_inches=6.5, caption="Hình 5 — Selected 196 đơn READY, bấm Create Batch")

add_para(doc, "Lưu ý:", bold=True)
add_bullet(doc, "Batch là MỘT CHIỀU: tạo rồi không bỏ được. Cẩn thận chọn đúng đơn trước khi bấm.")
add_bullet(doc, "1 batch có thể chứa đơn của nhiều khách + nhiều sản phẩm.")
add_bullet(doc, "Số thứ tự batch (-001, -002, -003) auto increment trong ngày, theo AM/PM.")

doc.add_page_break()

# --- STEP 5: DOWNLOAD FILE ---
add_heading(doc, "Bước 5 — DOWNLOAD FILE: Tải file Excel batch để up ClickShip/EST", level=2)

add_para(doc, "Mục đích:", bold=True)
add_para(
    doc,
    "Tải file Excel chứa toàn bộ đơn của batch (đúng format ClickShip yêu cầu) để up lên hệ "
    "thống của nhà cung cấp vận chuyển tạo label.",
)

add_para(doc, "Thao tác:", bold=True)
add_bullet(doc, "Sidebar → Batches.")
add_bullet(doc, "Trang Batches list tất cả batch đã tạo, sắp xếp theo ngày tạo mới nhất.")
add_bullet(doc, 'Tìm batch vừa tạo (vd 2026-05-27-PM-003) → bấm "Download Excel" cột Actions.')
add_bullet(doc, "File .xlsx download xuống máy.")
add_bullet(doc, "Up file này vào ClickShip (hoặc EST) → ClickShip tạo label cho từng đơn.")

add_image_centered(doc, IMG["download"], width_inches=6.5, caption="Hình 6 — Trang Batches, download Excel batch 2026-05-27-PM-003")

add_para(doc, "Lưu ý:", bold=True)
add_bullet(doc, "File Excel có đầy đủ cột ClickShip cần: #TITLENAME, Name, #LASTNAME, Address fields, Phone, Quantity, Payment method, COD amount...")
add_bullet(doc, "Nếu khách thiếu công thức cột nào trong sheet → app sẽ pull về DB với field đó empty → file Excel sẽ thiếu cột đó.")
add_bullet(doc, "Có thể download lại batch nhiều lần — file luôn được generate live từ DB.")

doc.add_page_break()

# --- STEP 6: UPLOAD LABELED ---
add_heading(doc, "Bước 6 — UPLOAD LABELED: Đối soát file kết quả từ ClickShip/EST", level=2)

add_para(doc, "Mục đích:", bold=True)
add_para(
    doc,
    "Sau khi ClickShip tạo label xong, download file kết quả (Order Fulfilled) từ ClickShip "
    "và up lại vào app. App sẽ match orderId → cập nhật tracking number, URL, carrier, "
    "ship date vào DB → đẩy ngược về sheet của khách (tự động).",
)

add_para(doc, "Thao tác:", bold=True)
add_bullet(doc, "Trên ClickShip: download file Order Fulfilled (theo ngày — file sẽ chứa toàn bộ đơn fulfilled trong ngày, không chỉ riêng batch của mình).")
add_bullet(doc, "Trong app: Sidebar → Carrier Tracking → tab Create Label.")
add_bullet(doc, "Chọn batch tương ứng từ dropdown (vd 2026-05-27-PM-003).")
add_bullet(doc, 'Chọn file Excel vừa download → bấm "Import Labels".')
add_bullet(doc, "App match orderId → cập nhật tracking → status: EXPORTED → LABEL_CREATED.")
add_bullet(doc, "App tự động đẩy tracking ngược về Google Sheet của khách (sheet writeback).")

add_image_centered(doc, IMG["upload"], width_inches=6.5, caption="Hình 7 — Shipping Reconciliation, chọn batch + upload file ClickShip")

add_para(doc, "Lưu ý quan trọng:", bold=True)
add_bullet(doc, "Đơn nào trong batch CÓ trong file → LABEL_CREATED + có tracking.")
add_bullet(doc, "Đơn nào trong batch KHÔNG có trong file → ERROR (ClickShip không tạo được label, có thể do data lỗi địa chỉ/zipcode).")
add_bullet(doc, "Đơn ERROR có thể sửa data + tạo label lại trong ClickShip, sau đó up file mới → app cập nhật ERROR → LABEL_CREATED.")

doc.add_page_break()

# --- STEP 7: ERROR HANDLING ---
add_heading(doc, "Bước 7 — XỬ LÝ ĐƠN ERROR: Báo khách lỗi để fix", level=2)

add_para(doc, "Mục đích:", bold=True)
add_para(
    doc,
    "Đơn ERROR có 2 nguồn:",
)
add_bullet(doc, "Lỗi data lúc parse từ sheet: thiếu Name / Address / Phone (vd Phone bị #ERROR! do công thức sai).", bold_first="(a) ")
add_bullet(doc, 'Lỗi ClickShip từ chối tạo label: thường do zipcode/province sai format (vd zipcode có dấu chấm thừa, chữ "O" thay vì số "0", province "British Columbia" thay vì "BC").', bold_first="(b) ")

add_para(doc, "Cách kiểm tra:", bold=True)
add_bullet(doc, "Vào Active Orders → bấm filter chip Error (chấm đỏ) phía trên bảng.")
add_bullet(doc, "Bảng hiển thị danh sách đơn ERROR + lý do (cột errorNote).")
add_bullet(doc, "Click vào đơn để xem chi tiết.")

add_image_centered(doc, IMG["error"], width_inches=6.5, caption="Hình 8 — Active Orders filter Error, 2 đơn cần báo khách sửa")

add_para(doc, "Quy trình xử lý:", bold=True)
add_bullet(doc, "Identify lý do lỗi (data thiếu vs ClickShip reject).")
add_bullet(doc, "Báo khách qua kênh phù hợp (Zalo/Email/WhatsApp) để fix data trên Google Sheet.")
add_bullet(doc, "Khách fix xong → sync lại (Bước 1) → đơn chuyển ERROR_UPDATED → validate lại → READY.")
add_bullet(doc, "Đưa lại vào batch (Bước 3-5) → up ClickShip → tracking về (Bước 6).")

add_para(doc, "Mẹo:", bold=True)
add_bullet(doc, "Có thể gắn cờ đơn (red/yellow flag) để chat 2 chiều KDE ↔ khách qua page Flagged.")
add_bullet(doc, "Đơn ERROR đã có batch + đã sửa: app tự nhận khi up file ClickShip mới (sau commit e00c91d, 2026-05-27).")

doc.add_page_break()

# =================================================================
# III. CHECKLIST
# =================================================================
add_heading(doc, "III. Checklist vận hành hàng ngày", level=1)

add_heading(doc, "Buổi sáng (AM)", level=2)
add_bullet(doc, "Dashboard → bấm Sync để pull đơn mới từ tất cả sheet khách.")
add_bullet(doc, "Kiểm tra READY TO PACK count.")
add_bullet(doc, "Kiểm tra NEEDS ATTENTION (đơn delay, stuck, notice card).")
add_bullet(doc, "Kiểm tra RED/YELLOW FLAG orders (khách có gắn cờ đợi trả lời không).")
add_bullet(doc, "Filter Error → báo khách lỗi nếu có đơn mới ERROR.")

add_heading(doc, "Lúc tạo batch", level=2)
add_bullet(doc, "Active Orders → filter Ready.")
add_bullet(doc, "Chọn đơn → Create Batch.")
add_bullet(doc, "Batches → Download Excel → up ClickShip/EST.")

add_heading(doc, "Sau khi ClickShip xong", level=2)
add_bullet(doc, "Download file Order Fulfilled.")
add_bullet(doc, "Carrier Tracking → chọn batch → Import Labels.")
add_bullet(doc, "Kiểm tra số đơn ERROR (ClickShip không tạo được label) → báo khách.")
add_bullet(doc, "Verify trên Google Sheet khách đã có tracking number (auto-sync ngay sau import).")

add_heading(doc, "Cuối ngày", level=2)
add_bullet(doc, "Kiểm tra Dashboard: tỷ lệ DELIVERED, đơn IN TRANSIT bao nhiêu.")
add_bullet(doc, "Trả lời các cờ khách nhắn trong ngày.")

# Save
doc.save(OUT_PATH)
print(f"Generated: {OUT_PATH}")
