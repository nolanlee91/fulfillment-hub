# -*- coding: utf-8 -*-
"""Build customer guide docx for Fulfillment Hub (CosmeticsTuanlx)."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = os.path.join("Hướng dẫn khách hàng", "1. CosmestisTuanlx")
CEN = os.path.join("scripts", "_censored")
OUT = os.path.join("Hướng dẫn khách hàng",
                   "Huong dan su dung Fulfillment Hub - CosmeticsTuanlx.docx")

# (heading, body, image_path)
o = lambda f: os.path.join(SRC, f)
c = lambda k: os.path.join(CEN, k + ".png")

STEPS = [
    ("1. Theo dõi các đơn đang xử lý (Active Orders)",
     "Đăng nhập vào app, chọn mục Active Orders ở menu bên trái. Đây là danh sách "
     "tất cả đơn hàng đang được xử lý. Thanh trạng thái phía trên (New, Ready, Error, "
     "Exported, Label, In Transit) giúp lọc nhanh từng nhóm. Dòng “X need attention” "
     "cho biết số đơn đang cần bạn kiểm tra.",
     o("1. Các đơn hàng đang active.png")),

    ("2. Lọc các đơn bị lỗi (Error)",
     "Bấm vào tab Error trên thanh trạng thái để chỉ hiển thị những đơn đang bị lỗi "
     "(thiếu hoặc sai thông tin). Đây là các đơn cần bạn cập nhật lại trên Google Sheet.",
     c("2")),

    ("3. Xem đơn lỗi ở đâu để sửa trên Google Sheet",
     "Bấm vào một đơn để mở bảng chi tiết bên phải. Mục DATA ERROR (ví dụ "
     "“Missing: Phone”) cho biết đơn đang thiếu hoặc sai thông tin gì. Bạn quay lại "
     "Google Sheet, bổ sung đúng thông tin đó — hệ thống sẽ tự đồng bộ và đơn sẽ hết lỗi.",
     c("3")),

    ("3.1. Lọc nhanh theo loại lỗi (Attention)",
     "Dùng bộ lọc Attention để lọc nhanh các đơn theo loại lỗi / cờ hay gặp, giúp bạn "
     "xử lý hàng loạt nhanh hơn.",
     c("31")),

    ("4. Gắn cờ đơn cần KDExpress hỗ trợ",
     "Với những đơn cần KDExpress hỗ trợ, bạn gắn cờ cho đơn. Các đơn đã gắn cờ nằm ở "
     "mục Flagged — nơi bạn và KDExpress trao đổi qua lại (Red Flag / Yellow Flag / "
     "Resolved).",
     o("4. Gắn cờ các đơn cần KDexpress hỗ trợ.png")),

    ("5. Xem các đơn đã giao thành công (Delivered)",
     "Vào mục Delivered để xem những đơn đã giao thành công. Có thể lọc theo sản phẩm, "
     "hình thức thanh toán, hoặc tìm kiếm; bấm Export để tải danh sách về.",
     o("5. Xem các đơn đã giao thành công.png")),

    ("6. Xem các đơn giao thất bại / hoàn về (Failed)",
     "Vào mục Failed để xem các đơn giao không thành công hoặc bị hoàn về người gửi "
     "(returned to sender).",
     o("6. Xem các đơn giao hàng thất bại.png")),

    ("7. Đối soát — Tải file mẫu (Template)",
     "Vào mục Reconciliation. Với thanh toán chuyển khoản ETF, bấm Download template để "
     "tải file Excel mẫu gồm 2 cột: Order ID và Ref Number.",
     o("7. Đối soát download tempalte.png")),

    ("7.1. Điền file đối soát rồi tải lên app",
     "Điền 2 cột vào file: Order ID (lấy ở Google Sheet) và Ref Number (mã tham chiếu "
     "trong email báo nhận tiền). Lưu file lại rồi tải lên app ở mục Reconciliation — "
     "hệ thống sẽ tự khớp tiền vào từng đơn.",
     o("7.1 Làm file đối soát, sau đó lưu lại và update lên app. Mã order ở google sheet, mã Ref ở email thông báo nhận tiền.png")),

    ("7.2. Trường hợp KHÔNG phải ETF",
     "Nếu khoản thanh toán không phải ETF (ví dụ bank transfer, money order, cheque…), "
     "bạn mở đúng đơn hàng đó, tại mục NON-ETF bấm Choose file để chọn ảnh chứng từ, "
     "rồi bấm Upload.",
     c("72")),
]

doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# margins
for s in doc.sections:
    s.left_margin = Inches(0.8)
    s.right_margin = Inches(0.8)
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)

CONTENT_W = Inches(7.0)
GREEN = RGBColor(0x1F, 0x7A, 0x33)

# title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("HƯỚNG DẪN SỬ DỤNG FULFILLMENT HUB")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = GREEN

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("KDExpress – Dành cho khách hàng")
rs.font.size = Pt(12)
rs.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

intro = doc.add_paragraph(
    "Tài liệu hướng dẫn các thao tác cơ bản trên app quản lý đơn hàng: theo dõi đơn, "
    "xử lý đơn lỗi, gắn cờ nhờ hỗ trợ, xem đơn đã giao / thất bại và đối soát thanh toán."
)
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

for i, (head, body, img) in enumerate(STEPS):
    h = doc.add_heading(head, level=1)
    for run in h.runs:
        run.font.color.rgb = GREEN
    doc.add_paragraph(body)
    doc.add_picture(img, width=CONTENT_W)
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if i != len(STEPS) - 1:
        doc.add_paragraph()

doc.save(OUT)
print("Saved:", OUT)
